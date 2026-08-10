import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = "/home/ubuntu/research_note.md"
OUT = "/home/ubuntu/Aaapurti_SEZ_FTWZ_SIIB_Research_Note.docx"

doc = Document()
st = doc.styles["Normal"]
st.font.name = "Times New Roman"
st.font.size = Pt(11)
for s in doc.sections:
    s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Inches(1)


def add_runs(p, text):
    for part in re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            p.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            p.add_run(part[1:-1]).italic = True
        elif part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1])
            r.font.name = "Courier New"
        else:
            p.add_run(part)


lines = open(SRC, encoding="utf-8").read().split("\n")
i = 0
while i < len(lines):
    line = lines[i].rstrip()
    if line.strip() == "---":
        i += 1
        continue
    if line.startswith("|") and i + 1 < len(lines) and set(lines[i + 1].replace("|", "").replace(" ", "")) <= set("-:"):
        rows = []
        header = [c.strip() for c in line.strip("|").split("|")]
        i += 2
        while i < len(lines) and lines[i].startswith("|"):
            rows.append([c.strip() for c in lines[i].strip("|").split("|")])
            i += 1
        t = doc.add_table(rows=1, cols=len(header))
        t.style = "Table Grid"
        for j, h in enumerate(header):
            cell = t.rows[0].cells[j]
            cell.text = ""
            add_runs(cell.paragraphs[0], h)
            for r in cell.paragraphs[0].runs:
                r.bold = True
        for row in rows:
            cells = t.add_row().cells
            for j, c in enumerate(row[: len(header)]):
                cells[j].text = ""
                add_runs(cells[j].paragraphs[0], c)
        doc.add_paragraph()
        continue
    m = re.match(r"^(#{1,6})\s+(.*)$", line)
    if m:
        lvl = min(len(m.group(1)), 4)
        p = doc.add_heading(level=lvl)
        add_runs(p, m.group(2))
        i += 1
        continue
    if re.match(r"^\s*[-*]\s+", line):
        p = doc.add_paragraph(style="List Bullet")
        add_runs(p, re.sub(r"^\s*[-*]\s+", "", line))
        i += 1
        continue
    if re.match(r"^\s*\d+\.\s+", line):
        p = doc.add_paragraph(style="List Number")
        add_runs(p, re.sub(r"^\s*\d+\.\s+", "", line))
        i += 1
        continue
    if line.startswith(">"):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        add_runs(p, line.lstrip("> ").strip())
        for r in p.runs:
            r.italic = True
        i += 1
        continue
    if not line.strip():
        i += 1
        continue
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_runs(p, line)
    i += 1

doc.save(OUT)
print("ok")
