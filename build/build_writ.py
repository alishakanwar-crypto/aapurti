import sys
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK

ROWS = [
    ("1", "5917422", "25.11.2025", "10.154", "16,110.00"),
    ("2", "6496054", "24.12.2025", "69.032", "1,11,104.00"),
    ("3", "6558975", "27.12.2025", "101.032", "1,61,874.00"),
    ("4", "6836417", "12.01.2026", "101.115", "1,64,348.00"),
    ("5", "6836416", "12.01.2026", "10.180", "16,525.60"),
    ("6", "6879260", "14.01.2026", "1.988", "4,03,179.24"),
    ("7", "7115035", "25.01.2026", "105.372", "22,51,837.44"),
    ("8", "7173539", "29.01.2026", "312.870", "2,12,40,558.43"),
    ("9", "8371333", "30.03.2026", "101.483", "2,08,74,925.44"),
    ("10", "9008650", "01.05.2026", "132.137", "2,71,51,058.76"),
    ("11", "9926786", "17.06.2026", "305.518", "6,06,61,838.24"),
    ("12", "2202151", "30.06.2026", "465.992", "8,92,15,669.66"),
    ("13", "2321952", "05.07.2026", "342.181", "6,50,06,626.94"),
    ("14", "2461736", "12.07.2026", "330.834", "6,64,83,317.94"),
    ("15", "2609344", "19.07.2026", "472.130", "9,60,84,218.84"),
    ("16", "2752795", "26.07.2026", "420.813", "8,54,83,964.95"),
    ("17", "2752867", "26.07.2026", "404.658", "8,15,22,804.46"),
    ("18", "2875203", "01.08.2026", "395.276", "7,78,55,964.59"),
]

doc = Document()
st = doc.styles["Normal"]
st.font.name = "Times New Roman"
st.font.size = Pt(12)
st.paragraph_format.line_spacing = 1.5
st.paragraph_format.space_after = Pt(6)
for s in doc.sections:
    s.left_margin = Inches(1.5)
    s.right_margin = Inches(1.0)


def add(text="", align=None, bold=False, indent=None, first=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    if indent is not None:
        p.paragraph_format.left_indent = Inches(indent)
    if first is not None:
        p.paragraph_format.first_line_indent = Inches(first)
    r = p.add_run(text)
    r.bold = bold
    return p


counters = {"N": 0, "GN": 0, "DN": 0}
cur_table = None
cur_widths = None
cur_serial = 0

SRC = sys.argv[1] if len(sys.argv) > 1 else "/home/ubuntu/writ_draft.txt"
OUT = sys.argv[2] if len(sys.argv) > 2 else "/home/ubuntu/Aaapurti_Writ_Petition_Madras_HC_DRAFT.docx"

with open(SRC) as f:
    lines = f.read().split("\n")

for line in lines:
    if not line.strip():
        continue
    tag, _, text = line.partition("|")
    if tag == "BLANK":
        doc.add_paragraph()
    elif tag == "PB":
        doc.add_page_break()
        counters["N"] = 0
    elif tag == "H1":
        add(text, WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    elif tag == "C":
        add(text, WD_ALIGN_PARAGRAPH.CENTER)
    elif tag == "B":
        add(text, WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    elif tag == "P":
        add(text, WD_ALIGN_PARAGRAPH.JUSTIFY)
    elif tag in ("N", "GN", "DN"):
        counters[tag] += 1
        add(f"{counters[tag]}.  {text}", WD_ALIGN_PARAGRAPH.JUSTIFY, indent=0.4, first=-0.4)
    elif tag == "TSTART":
        headers = text.split(";;")
        widths = None
        if "@@" in headers[-1]:
            headers[-1], w = headers[-1].split("@@")
            widths = [float(x) for x in w.split(",")]
        cur_table = doc.add_table(rows=1, cols=len(headers))
        cur_table.style = "Table Grid"
        cur_table.autofit = False
        for i, h in enumerate(headers):
            c = cur_table.rows[0].cells[i]
            c.text = ""
            run = c.paragraphs[0].add_run(h)
            run.bold = True
            run.font.size = Pt(11)
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            c.paragraphs[0].paragraph_format.line_spacing = 1.0
            c.paragraphs[0].paragraph_format.space_after = Pt(0)
        cur_widths = widths
        cur_serial = 0
    elif tag == "TROW":
        vals = text.split(";;")
        cur_serial += 1
        vals = [str(cur_serial) + "."] + vals
        cells = cur_table.add_row().cells
        for i, v in enumerate(vals):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(0)
            if i in (0, 1, len(vals) - 1):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(v)
            r.font.size = Pt(11)
    elif tag == "TEND":
        if cur_widths:
            for row in cur_table.rows:
                for i, w in enumerate(cur_widths):
                    row.cells[i].width = Inches(w)
        doc.add_paragraph()
        cur_table = None
        cur_widths = None
    elif tag == "TABLE":
        t = doc.add_table(rows=1, cols=5)
        t.style = "Table Grid"
        hdr = ["S. No.", "Bill of Entry No.", "Date", "Quantity (kg)", "Value (Rs.)"]
        for i, h in enumerate(hdr):
            c = t.rows[0].cells[i]
            c.text = ""
            run = c.paragraphs[0].add_run(h)
            run.bold = True
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for row in ROWS:
            cells = t.add_row().cells
            for i, v in enumerate(row):
                cells[i].text = v
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    p.paragraph_format.line_spacing = 1.0
                    p.paragraph_format.space_after = Pt(0)
                    for run in p.runs:
                        run.font.size = Pt(10)
        doc.add_paragraph()

doc.save(OUT)
print("ok")
